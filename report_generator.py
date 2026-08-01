"""Template-faithful DOCX generation for the RSS Plan Builder."""

from __future__ import annotations

import hashlib
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TEMPLATE_PATH = (
    Path(__file__).resolve().parent
    / "assets"
    / "Remote Site Supervision Plan (Template).docx"
)
TEMPLATE_SHA256 = (
    "1A9DCE5E3006361CE001DDB97143BA58B6734D773877DAE53A1B00B8184F074B"
)
BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(96, 96, 96)


def _text(value: Any, blank: str = "") -> str:
    value = "" if value is None else str(value).strip()
    return value or blank


def _verify_template(path: Path) -> None:
    if not path.exists():
        raise FileNotFoundError(
            "The BCA Remote Site Supervision Plan template is missing from "
            f"{path.parent}."
        )
    digest = hashlib.sha256(path.read_bytes()).hexdigest().upper()
    if path == TEMPLATE_PATH and digest != TEMPLATE_SHA256:
        raise ValueError(
            "The bundled BCA template has changed. Restore the original template "
            "before generating a report."
        )


def _set_run_style(run, size: float = 10, bold: bool = False, italic: bool = False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK


def _set_paragraph_style(
    paragraph,
    size: float = 10,
    bold: bool = False,
    after: float = 4,
    before: float = 0,
):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.05
    for run in paragraph.runs:
        _set_run_style(run, size=size, bold=bold)


def _clear_cell(cell) -> None:
    cell._element.clear_content()
    cell.add_paragraph()


def _write_cell(cell, value: Any, size: float = 10, bold: bool = False) -> None:
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(_text(value))
    _set_run_style(run, size=size, bold=bold)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _write_labeled_cell(cell, label: str, value: Any) -> None:
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    label_run = paragraph.add_run(f"{label}:\n")
    _set_run_style(label_run, size=10, bold=True)
    value_run = paragraph.add_run(_text(value, "Not specified"))
    _set_run_style(value_run, size=10)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_label_value(cell, label: str, value: Any, size: float = 10) -> None:
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.05
    label_run = paragraph.add_run(f"{label}: ")
    _set_run_style(label_run, size=size, bold=True)
    value_run = paragraph.add_run(_text(value, "Not specified"))
    _set_run_style(value_run, size=size)


def _add_subheading(cell, text: str) -> None:
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    _set_run_style(run, size=11, bold=True)


def _set_cell_margins(
    cell, top: int = 100, start: int = 110, bottom: int = 100, end: int = 110
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _add_picture(
    cell,
    image_bytes: bytes | None,
    caption: str,
    placeholder: str,
    width: float = 5.55,
) -> None:
    if not image_bytes:
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(placeholder)
        _set_run_style(run, size=9, italic=True)
        run.font.color.rgb = GREY
        return
    try:
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(BytesIO(image_bytes), width=Inches(width))
        if caption:
            caption_paragraph = cell.add_paragraph()
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_paragraph.add_run(caption)
            _set_run_style(caption_run, size=9, italic=True)
    except Exception:
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(
            f"{placeholder} (The uploaded image could not be embedded.)"
        )
        _set_run_style(run, size=9, italic=True)
        run.font.color.rgb = GREY


def _add_two_column_table(
    container_cell, rows: list[tuple[str, Any]], label_width: float = 2.05
) -> None:
    table = container_cell.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total_width = 5.86
    for label, value in rows:
        cells = table.add_row().cells
        _prevent_row_split(table.rows[-1])
        cells[0].width = Inches(label_width)
        cells[1].width = Inches(total_width - label_width)
        for cell in cells:
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _write_cell(cells[0], label, size=9, bold=True)
        _write_cell(cells[1], _text(value, "Not specified"), size=9)


def _competency_summary(team: dict[str, Any]) -> str:
    checks = [
        (
            "Technology-provider training completed",
            team.get("competency_provider_training"),
        ),
        (
            "Proficiency demonstrated during an RSS trial",
            team.get("competency_trial"),
        ),
        (
            "Professional registration / appointment checked",
            team.get("competency_registration"),
        ),
        (
            "Upgrade / refresher training arrangement confirmed",
            team.get("competency_upgrade_training"),
        ),
    ]
    lines = [f"{'Confirmed' if status else 'Not confirmed'} - {label}" for label, status in checks]
    evidence = _text(team.get("competency_evidence") or team.get("competency"))
    if evidence:
        lines.append(f"Evidence / record reference - {evidence}")
    verifier = _text(team.get("competency_verifier"))
    date = _text(team.get("competency_date"))
    if verifier or date:
        lines.append(
            f"Verified by - {verifier or 'Not specified'}"
            f"; date - {date or 'Not specified'}"
        )
    return "\n".join(lines)


PROFILE_FIELDS = {
    "people": [
        ("Site supervisors", "site_supervisors"),
        ("Builder-side operators", "builder_operators"),
        ("Backup personnel", "backup_personnel"),
        ("Deployment / handover notes", "notes"),
    ],
    "technology": [
        ("Live-streaming devices", "live_devices"),
        ("Evidence / measurement devices", "evidence_devices"),
        ("Two-way audio", "audio"),
        ("Primary connectivity", "connectivity"),
        ("Backup connectivity", "backup_connectivity"),
        ("RSS platform / software", "platform"),
        ("Video / recording standard", "video_standard"),
        ("Secure storage", "storage"),
        ("Power / battery backup", "power_backup"),
        ("Equipment register / calibration", "equipment_register"),
    ],
    "controls": [
        ("Preparation", "before"),
        ("Live supervision", "during"),
        ("Close-out", "after"),
        ("Communication", "communication"),
        ("Stop-work / in-person trigger", "stop_work"),
        ("Technology failure", "tech_failure"),
        ("Poor evidence", "poor_evidence"),
        ("Safety incident", "safety_incident"),
        ("Non-conformity", "non_conformity"),
    ],
    "records": [
        ("Naming / indexing", "naming"),
        ("Access / data security", "access"),
        ("Backup / recovery", "backups"),
        ("Retention", "retention"),
        ("Verification", "verification"),
        ("Audits", "audits"),
        ("Performance monitoring", "performance"),
        ("Traceability", "traceability"),
    ],
}


def _profile_reference(data: dict[str, Any], kind: str, profile_id: Any) -> str:
    profile_id = _text(profile_id)
    for profile in data.get("profiles", {}).get(kind, []):
        if profile.get("id") == profile_id:
            return f"{profile_id} - {_text(profile.get('name'), 'Unnamed profile')}"
    return profile_id or "Not specified"


def _profile_details(profile: dict[str, Any], kind: str) -> str:
    if profile.get("default"):
        return (
            "Project default; uses the master information in the corresponding "
            "plan section."
        )
    variations = [
        f"{label}: {_text(profile.get(key))}"
        for label, key in PROFILE_FIELDS[kind]
        if _text(profile.get(key))
    ]
    return (
        "\n".join(variations)
        if variations
        else "No variation entered; inherits the project default."
    )


def _add_profile_register(
    cell, data: dict[str, Any], kind: str, heading: str
) -> None:
    profiles = data.get("profiles", {}).get(kind, [])
    if not profiles:
        return
    _add_subheading(cell, heading)
    _add_two_column_table(
        cell,
        [
            (
                f"{_text(profile.get('id'), 'Unnumbered')} - "
                f"{_text(profile.get('name'), 'Unnamed profile')}",
                _profile_details(profile, kind),
            )
            for profile in profiles
        ],
        label_width=2.05,
    )


def _phase_details(phase: dict[str, Any]) -> str:
    parts = [
        f"Activity progress range: {_text(phase.get('progress_range'), 'Not specified')}",
        f"Extent of RSS: {_text(phase.get('rss_extent'), 'Not specified')}",
        (
            "Parallel / in-person verification: "
            f"{_text(phase.get('parallel_supervision'), 'Use project default')}"
        ),
        (
            "Acceptance / progression criteria: "
            f"{_text(phase.get('acceptance_criteria'), 'Not specified')}"
        ),
        f"QP(S) review point: {_text(phase.get('review_point'), 'Not specified')}",
    ]
    if _text(phase.get("remarks")):
        parts.append(f"Remarks: {_text(phase.get('remarks'))}")
    return "\n".join(parts)


def _phase_name(phase: dict[str, Any], fallback: str) -> str:
    if phase.get("name") == "Custom":
        return _text(phase.get("custom_name"), "Custom phase")
    return _text(phase.get("name"), fallback)


def _fill_cover(doc: Document, data: dict[str, Any]) -> None:
    project = data.get("project", {})
    team = data.get("team", {})
    cover = doc.tables[0]
    _write_labeled_cell(cover.cell(0, 0), "Project Reference No.", project.get("reference"))
    _write_labeled_cell(cover.cell(1, 0), "Project Description", project.get("description"))
    _write_labeled_cell(
        cover.cell(2, 0),
        f"{_text(project.get('site_type'), 'Construction site')}\nAddress",
        project.get("address"),
    )

    prepared = doc.tables[1]
    _write_cell(prepared.cell(0, 0), "Name of Qualified Person (Supervision)\n& PE Registration No.", size=10)
    _write_cell(
        prepared.cell(0, 1),
        f"{_text(team.get('qp_name'), 'Not specified')}\n"
        f"PE Reg. No. {_text(team.get('pe_number'), 'Not specified')}",
        size=10,
    )
    _write_cell(prepared.cell(1, 0), "Company", size=10)
    _write_cell(prepared.cell(1, 1), team.get("company"), size=10)
    _write_cell(prepared.cell(2, 0), "Date", size=10)
    _write_cell(prepared.cell(2, 1), team.get("prepared_date"), size=10)


def _fill_project_background(
    doc: Document, data: dict[str, Any], site_plan_bytes: bytes | None
) -> None:
    project = data.get("project", {})
    cell = doc.tables[3].cell(0, 0)
    _clear_cell(cell)
    _add_label_value(cell, "Project description", project.get("description"))
    _add_label_value(cell, "Site location", project.get("address"))
    _add_label_value(cell, "Structural system", project.get("structural_system"))
    _add_label_value(cell, "Foundation system", project.get("foundation_system"))
    _add_label_value(cell, "RSS constraints / challenges", project.get("challenges"))
    _add_label_value(cell, "Permit date", project.get("permit_date"))

    site_table = doc.tables[4]
    _clear_cell(site_table.cell(0, 0))
    _add_picture(
        site_table.cell(0, 0),
        site_plan_bytes,
        "",
        "Insert overall site plan or location of construction site / fabrication yard.",
    )
    _write_cell(
        site_table.cell(1, 0),
        "Figure 1: Overall Site Plan or Location of Fabrication Yard",
        size=9,
    )


def _fill_phasing(doc: Document, data: dict[str, Any]) -> None:
    phases = data.get("phases", {})
    cell = doc.tables[5].cell(0, 0)
    _clear_cell(cell)
    _add_subheading(cell, "2.1. Phased Implementation")
    _add_two_column_table(
        cell,
        [
            ("Phase 1 - first 30% of works", phases.get("phase_1")),
            ("Phase 2 - 30% to 75% of works", phases.get("phase_2")),
            ("Phase 3 - 75% to 100% of works", phases.get("phase_3")),
            ("Beyond Phase 3", phases.get("beyond")),
        ],
    )
    _add_subheading(cell, "2.2. Acceptance Criteria for RSS Implementation")
    _add_label_value(cell, "Phase acceptance criteria", phases.get("criteria"))
    _add_label_value(cell, "Parallel supervision plan", phases.get("parallel_plan"))
    _add_label_value(cell, "Review / adjustment cadence", phases.get("review_cadence"))


def _fill_people(
    doc: Document, data: dict[str, Any], org_chart_bytes: bytes | None
) -> None:
    team = data.get("team", {})
    cell = doc.tables[6].cell(0, 0)
    _clear_cell(cell)
    _add_subheading(cell, "3.1. RSS Team Structure")
    _add_label_value(cell, "Organisation and reporting lines", team.get("organisation"))
    _add_picture(
        cell,
        org_chart_bytes,
        "Figure 2: RSS Team Organisation and Reporting Lines",
        "Optional: insert organisation / reporting-line chart.",
        width=5.2,
    )
    _add_subheading(cell, "3.2. Roles and Responsibilities")
    _add_label_value(cell, "Qualified Person (Supervision)", team.get("qp_name"))
    _add_label_value(cell, "Site supervisors (RE/RTO)", team.get("site_supervisors"))
    _add_label_value(cell, "Builder-side RSS operators", team.get("builder_operators"))
    _add_label_value(cell, "Backup personnel and handover", team.get("backup_personnel"))
    _add_subheading(cell, "3.3. Training and Competency Requirements")
    _add_label_value(cell, "Training programme", team.get("training"))
    _add_label_value(cell, "Competency verification", _competency_summary(team))
    _add_profile_register(
        cell,
        data,
        "people",
        "3.4. Reusable Personnel Deployment Profiles",
    )


def _fill_activities(doc: Document, data: dict[str, Any]) -> None:
    activities = data.get("activities") or []
    table = doc.tables[7]
    note_cell = table.cell(0, 0)
    _clear_cell(note_cell)
    note = note_cell.paragraphs[0]
    note_run = note.add_run(
        "Each activity below records the QP(S)'s assessment, selected supervision "
        "approach and minimum evidence. The applicable checklist from the older "
        "Site Supervision Plan guide and Annex D of the RSS Guidebook must be "
        "reviewed for the project-specific scope."
    )
    _set_run_style(note_run, size=9)
    note.paragraph_format.space_after = Pt(0)

    cell = table.cell(1, 0)
    _clear_cell(cell)
    if not activities:
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run("No structural supervision activity has been selected.")
        _set_run_style(run, size=10, italic=True)
        return
    for index, activity in enumerate(activities, start=1):
        _add_subheading(
            cell,
            f"Activity {index}: {_text(activity.get('work_type'), 'Not selected')}",
        )
        _add_two_column_table(
            cell,
            [
                ("Guide category", activity.get("category")),
                ("Location / element IDs", activity.get("location")),
                ("Scope of supervision", activity.get("description")),
                (
                    "Assessment",
                    f"{_text(activity.get('complexity'), 'Not specified')} inspection / "
                    f"{_text(activity.get('frequency'), 'Not specified')} supervision",
                ),
                ("Selected approach", activity.get("approach")),
                (
                    "Assigned people profile",
                    _profile_reference(
                        data, "people", activity.get("people_profile_id")
                    ),
                ),
                (
                    "Assigned technology profile",
                    _profile_reference(
                        data,
                        "technology",
                        activity.get("technology_profile_id"),
                    ),
                ),
                (
                    "Assigned control profile",
                    _profile_reference(
                        data, "controls", activity.get("control_profile_id")
                    ),
                ),
                (
                    "Assigned record profile",
                    _profile_reference(
                        data, "records", activity.get("record_profile_id")
                    ),
                ),
                (
                    "Personnel requirements / variation",
                    activity.get("personnel_requirements"),
                ),
                ("Evidence requirements", activity.get("evidence")),
                ("Equipment / software variation", activity.get("equipment")),
                (
                    "Activity-specific controls / hold points",
                    activity.get("control_overrides"),
                ),
                (
                    "Record / retention / verification variation",
                    activity.get("record_overrides"),
                ),
                (
                    "Guide / Annex D review",
                    "Confirmed"
                    if activity.get("annex_d_reviewed")
                    else "Not confirmed",
                ),
                ("Professional justification", activity.get("deviation")),
            ],
            label_width=1.8,
        )
        phases = activity.get("implementation_phases") or []
        if not phases and _text(activity.get("phase")):
            phases = [
                {
                    "name": activity.get("phase"),
                    "progress_range": "Not specified",
                    "rss_extent": activity.get("extent"),
                    "parallel_supervision": "Use project default",
                    "acceptance_criteria": "Use project default",
                    "review_point": "Use project default",
                }
            ]
        _add_subheading(cell, f"Implementation Phases - Activity {index}")
        if phases:
            _add_two_column_table(
                cell,
                [
                    (
                        _phase_name(phase, f"Phase entry {phase_index}"),
                        _phase_details(phase),
                    )
                    for phase_index, phase in enumerate(phases, start=1)
                ],
                label_width=1.5,
            )
        else:
            _add_label_value(cell, "Implementation phases", "Not specified")


def _fill_devices_and_infrastructure(doc: Document, data: dict[str, Any]) -> None:
    technology = data.get("technology", {})
    devices = doc.tables[8].cell(0, 0)
    _clear_cell(devices)
    for label, key in (
        ("Live streaming devices", "live_devices"),
        ("Evidence / measurement devices", "evidence_devices"),
        ("Two-way audio", "audio"),
        ("Power / battery backup", "power_backup"),
        ("Equipment register / calibration", "equipment_register"),
    ):
        _add_label_value(devices, label, technology.get(key))

    infrastructure = doc.tables[9].cell(0, 0)
    _clear_cell(infrastructure)
    for label, key in (
        ("Primary connectivity", "connectivity"),
        ("Backup connectivity", "backup_connectivity"),
        ("RSS platform / software", "platform"),
        ("Video / recording standard", "video_standard"),
        ("Secure storage", "storage"),
    ):
        _add_label_value(infrastructure, label, technology.get(key))
    _add_profile_register(
        infrastructure,
        data,
        "technology",
        "Reusable Technology Profiles",
    )


def _fill_process_quality_records(doc: Document, data: dict[str, Any]) -> None:
    process = data.get("process", {})
    records = data.get("records", {})
    process_cell = doc.tables[10].cell(0, 0)
    _clear_cell(process_cell)
    for label, key in (
        ("Before Remote Supervision (Preparation Works)", "before"),
        ("During Remote Supervision", "during"),
        ("After Remote Supervision", "after"),
        ("Communication Protocol", "communication"),
    ):
        _add_subheading(process_cell, label)
        paragraph = process_cell.add_paragraph()
        run = paragraph.add_run(_text(process.get(key), "Not specified"))
        _set_run_style(run, size=10)
        paragraph.paragraph_format.space_after = Pt(4)
    _add_profile_register(
        process_cell,
        data,
        "controls",
        "Reusable Control Profiles",
    )

    quality = doc.tables[11].cell(0, 0)
    _clear_cell(quality)
    for label, value in (
        ("Stop-work / in-person trigger", process.get("stop_work")),
        ("Technology failure", process.get("tech_failure")),
        ("Poor or incomplete evidence", process.get("poor_evidence")),
        ("Safety incident", process.get("safety_incident")),
        ("Non-conformity and escalation", process.get("non_conformity")),
        ("Verification plan", records.get("verification")),
        ("Audits / management review", records.get("audits")),
        ("Performance monitoring", records.get("performance")),
    ):
        _add_label_value(quality, label, value)

    record_cell = doc.tables[12].cell(0, 0)
    _clear_cell(record_cell)
    for label, key in (
        ("Naming / indexing", "naming"),
        ("Access and data security", "access"),
        ("Backup and recovery", "backups"),
        ("Retention schedule", "retention"),
        ("Session traceability", "traceability"),
    ):
        _add_label_value(record_cell, label, records.get(key))
    _add_profile_register(
        record_cell,
        data,
        "records",
        "Reusable Record Profiles",
    )

    signoff = data.get("signoff", {})
    team = data.get("team", {})
    _add_subheading(record_cell, "QP(S) Declaration")
    declaration = record_cell.add_paragraph()
    declaration_text = (
        "I confirm that this Remote Site Supervision Plan has been prepared "
        "using professional judgement for the project and activities described; "
        "the applicable supervision checklists and minimum requirements have "
        "been reviewed; in-person supervision will be used whenever the intended "
        "supervision outcome cannot be achieved remotely; and records will be "
        "controlled and retained in accordance with this plan and prevailing "
        "requirements."
    )
    declaration_run = declaration.add_run(declaration_text)
    _set_run_style(declaration_run, size=10)
    _add_label_value(
        record_cell,
        "Qualified Person (Supervision)",
        f"{_text(signoff.get('qp_signature'), 'Not specified')} / "
        f"PE Reg. No. {_text(team.get('pe_number'), 'Not specified')}",
    )
    _add_label_value(record_cell, "Date", signoff.get("sign_date"))


def _set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def build_docx(
    data: dict[str, Any],
    site_plan_bytes: bytes | None = None,
    org_chart_bytes: bytes | None = None,
    template_path: str | Path | None = None,
) -> bytes:
    """Fill a copy of the supplied BCA template and return a native DOCX."""
    path = Path(template_path) if template_path else TEMPLATE_PATH
    _verify_template(path)
    doc = Document(path)

    _fill_cover(doc, data)
    _fill_project_background(doc, data, site_plan_bytes)
    _fill_phasing(doc, data)
    _fill_people(doc, data, org_chart_bytes)
    _fill_activities(doc, data)
    _fill_devices_and_infrastructure(doc, data)
    _fill_process_quality_records(doc, data)
    _set_update_fields(doc)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()
