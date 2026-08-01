import io
import unittest
from zipfile import ZipFile

from docx import Document

from report_generator import build_docx


class ReportGeneratorTest(unittest.TestCase):
    def test_generates_valid_docx_with_expected_sections(self):
        data = {
            "project": {
                "reference": "A1234-TEST",
                "description": "Worked example",
                "site_type": "Construction site",
                "address": "1 Example Road, Singapore",
                "structural_system": "Reinforced-concrete frame",
                "foundation_system": "Bored piles",
                "challenges": "None identified",
                "permit_date": "2026-07-01",
            },
            "team": {
                "qp_name": "Er Example",
                "pe_number": "12345",
                "company": "Example Consultants",
                "prepared_date": "2026-08-01",
                "organisation": "QP(S) → RE → RTO → Builder operator",
                "site_supervisors": "Er Resident (RE)",
                "builder_operators": "Site operator",
                "backup_personnel": "Named trained backup",
                "training": "Provider and project RSS training completed",
                "competency": "Demonstrated during trial and accepted by QP(S)",
                "competency_provider_training": True,
                "competency_trial": True,
                "competency_registration": True,
                "competency_upgrade_training": True,
                "competency_evidence": "TR-001 and RSS-TRIAL-01",
                "competency_verifier": "Er Example, QP(S)",
                "competency_date": "2026-07-25",
            },
            "phases": {
                "phase_1": "15% of the first 30%",
                "phase_2": "30% from 30-75%",
                "phase_3": "50% from 75-100%",
                "beyond": "Not proposed",
                "criteria": "95% uptime and satisfactory performance",
                "parallel_plan": "10% sample and parallel learning phase",
                "review_cadence": "End of each phase",
            },
            "technology": {
                "live_devices": "Two HD cameras and one backup",
                "evidence_devices": "Digital calliper",
                "audio": "Noise-cancelling headset",
                "connectivity": "Primary 5G",
                "backup_connectivity": "Independent 4G router",
                "platform": "Controlled project platform",
                "video_standard": "1080p, timestamped",
                "storage": "Role-controlled project server",
                "power_backup": "Two spare batteries",
                "equipment_register": "Serialised and calibrated",
            },
            "process": {
                "before": "Verify drawings, method and equipment.",
                "during": "Inspect systematically and record decisions.",
                "after": "Save evidence and close actions.",
                "communication": "Repeat back critical instructions.",
                "stop_work": "Stop if visibility or intervention is inadequate.",
                "tech_failure": "Use backup, then revert in person.",
                "poor_evidence": "Reject and repeat.",
                "safety_incident": "Suspend and follow emergency plan.",
                "non_conformity": "Record, rectify and verify closure.",
            },
            "records": {
                "naming": "Project-Activity-Date-Revision",
                "access": "Role-based access",
                "backups": "Daily protected backup",
                "retention": "Reports 5 years; video 2 years after TOP",
                "verification": "10% in-person sample",
                "audits": "Planned internal audit",
                "performance": "Uptime and detection rate",
                "traceability": "Element, personnel, time, result and evidence",
            },
            "profiles": {
                "people": [
                    {
                        "id": "P1",
                        "name": "Project default RSS team",
                        "default": True,
                    },
                    {
                        "id": "P2",
                        "name": "Concrete deployment team",
                        "default": False,
                        "site_supervisors": "Er Resident",
                        "builder_operators": "Mr Operator",
                    },
                ],
                "technology": [
                    {
                        "id": "T1",
                        "name": "Project default technology set",
                        "default": True,
                    }
                ],
                "controls": [
                    {
                        "id": "C1",
                        "name": "Project default control profile",
                        "default": True,
                    }
                ],
                "records": [
                    {
                        "id": "R1",
                        "name": "Project default record profile",
                        "default": True,
                    }
                ],
            },
            "activities": [
                {
                    "work_type": "Concreting works",
                    "description": "Concrete placement supervision",
                    "location": "Block A, Level 2",
                    "complexity": "Simple",
                    "frequency": "Continuous",
                    "approach": "Live remote supervision",
                    "people_profile_id": "P2",
                    "technology_profile_id": "T1",
                    "control_profile_id": "C1",
                    "record_profile_id": "R1",
                    "implementation_phases": [
                        {
                            "name": "Phase 1",
                            "progress_range": "First 30% of the activity",
                            "rss_extent": "15% remote supervision",
                            "parallel_supervision": "10% parallel in-person sample",
                            "acceptance_criteria": "95% uptime",
                            "review_point": "QP(S) gate review",
                            "remarks": "",
                        },
                        {
                            "name": "Phase 2",
                            "progress_range": "30% to 75% of the activity",
                            "rss_extent": "30% remote supervision",
                            "parallel_supervision": "Targeted checks",
                            "acceptance_criteria": "No unresolved critical findings",
                            "review_point": "QP(S) gate review",
                            "remarks": "",
                        },
                    ],
                    "personnel_requirements": "Use P2 for every concrete pour.",
                    "evidence": "Full recording and identified screenshots",
                    "equipment": "",
                    "control_overrides": "Pre-pour hold point retained.",
                    "record_overrides": "",
                    "annex_d_reviewed": True,
                    "deviation": "",
                }
            ],
            "signoff": {
                "qp_signature": "Er Example",
                "sign_date": "2026-08-01",
            },
        }
        payload = build_docx(data)
        self.assertGreater(len(payload), 250_000)
        doc = Document(io.BytesIO(payload))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        with ZipFile(io.BytesIO(payload)) as package:
            document_xml = package.read("word/document.xml").decode("utf-8")
        self.assertIn("Remote Site", all_text)
        self.assertIn("PROJECT BACKGROUND", all_text)
        self.assertIn("A1234-TEST", table_text)
        self.assertIn("TR-001 and RSS-TRIAL-01", table_text)
        self.assertIn("Concreting works", table_text)
        self.assertIn("P2 - Concrete deployment team", document_xml)
        self.assertIn("30% to 75% of the activity", document_xml)
        self.assertIn("Pre-pour hold point retained.", document_xml)
        self.assertNotIn("Guide category", document_xml)
        self.assertGreaterEqual(len(doc.tables), 14)
        self.assertEqual(round(doc.sections[0].page_width.inches, 2), 8.27)
        self.assertEqual(round(doc.sections[0].page_height.inches, 2), 11.69)


if __name__ == "__main__":
    unittest.main()
